"""cover_letter.yaml (+ sibling instance.yaml for meta) -> cover_letter.pdf/.docx.

Phase 2 of the cover-letter track. This mirrors render_pdf.py deliberately:
the same LaTeX+Jinja2 delimiter remap, the same recursive tex-escape, the same
Tectonic compile and one-page check. Two properties are load-bearing:

  * The letterhead (name / contact / accent rule) is pulled from the sibling
    instance.yaml's `meta`, so the letter is guaranteed to match the resume
    that lives in the same output/<slug>/ folder — one source of truth for
    contact details, never re-typed.
  * The letter body is prose the tailor step (prompts/tailor_cover_letter.md,
    phase 3) already grounded in that same instance, so every claim here traces
    back to a bullet the resume actually selected from master.yaml.

cover_letter.yaml shape (produced by the tailor step; a minimal example ships
in the repo for render-only testing):

    schema_version: 1.0
    instance_ref: instance.yaml          # where meta + facts come from
    job_description_ref: job_description.txt
    date: "July 7, 2026"
    recipient:
      addressee: "Hiring Team"           # a named manager if the JD gives one
      company: "City of Toronto"
      location: "Toronto, ON"            # optional
    salutation: "Dear Hiring Team,"
    paragraphs:                          # 3-4 short paragraphs, one page
      - "<hook>"
      - "<proof>"
      - "<fit / enthusiasm>"
      - "<close + call to action>"
    signoff: "Sincerely,"
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Reuse the resume renderer's building blocks verbatim so the two pipelines
# can never drift on escaping rules or the Jinja/LaTeX delimiter remap.
from render_pdf import _escape_recursive, _jinja_env, compile_tex

import sys

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "templates"))
import resume_docx_style as style  # noqa: E402

TEMPLATE_NAME = "cover_letter.tex.j2"

# Letter-appropriate defaults: more breathing room than the resume's tightened
# 10pt / 0.4in-H / 0.3in-V. A ~300-word letter fits one page comfortably here.
DEFAULT_LETTER_LAYOUT = {
    "font_size_pt": 11.0,
    "margin_h_in": 1.0,
    "margin_v_in": 1.0,
}

REQUIRED_KEYS = ("date", "salutation", "paragraphs", "signoff")


def validate_letter(letter: dict) -> list[str]:
    """Cheap structural check (phase 2). The verbatim/truthfulness guard that
    ties paragraphs back to the instance lands with the tailor step in phase 3."""
    errors: list[str] = []
    for key in REQUIRED_KEYS:
        if not letter.get(key):
            errors.append(f"cover_letter.yaml missing required field: {key}")
    paras = letter.get("paragraphs")
    if isinstance(paras, list) and len(paras) > 5:
        errors.append(f"too many paragraphs ({len(paras)}); a one-page letter wants 3-4")
    return errors


def render_tex(letter: dict, meta: dict, out_dir: Path, layout: dict) -> Path:
    safe = _escape_recursive(letter)
    safe["meta"] = _escape_recursive(meta)
    safe["layout"] = layout

    env = _jinja_env()
    template = env.get_template(TEMPLATE_NAME)
    tex_source = template.render(**safe)

    out_dir.mkdir(parents=True, exist_ok=True)
    tex_path = out_dir / "cover_letter.tex"
    tex_path.write_text(tex_source, encoding="utf-8")
    return tex_path


def _extract_page_count(out_dir: Path) -> int | None:
    """Parse the lastpage macro out of cover_letter.aux (same technique as
    render_pdf.extract_page_count, keyed to this document's basename)."""
    aux_path = out_dir / "cover_letter.aux"
    if not aux_path.exists():
        return None
    aux_text = aux_path.read_text(encoding="utf-8", errors="replace")
    match = re.search(r"\\xdef\\lastpage@lastpage\{(\d+)\}", aux_text)
    return int(match.group(1)) if match else None


def render_pdf(letter: dict, meta: dict, out_dir: Path, layout: dict | None = None) -> dict:
    """Render cover_letter.tex, compile with Tectonic, extract page count.

    Returns {"success", "page_count", "pdf_path", "log"} — same shape as
    render_pdf.render_pdf so the CLI can treat both renders identically.
    """
    layout = {**DEFAULT_LETTER_LAYOUT, **(layout or {})}
    tex_path = render_tex(letter, meta, out_dir, layout)
    success, log = compile_tex(tex_path)
    pdf_path = out_dir / "cover_letter.pdf"
    page_count = _extract_page_count(out_dir) if success else None
    return {
        "success": success and pdf_path.exists(),
        "page_count": page_count,
        "pdf_path": pdf_path if pdf_path.exists() else None,
        "log": log,
    }


def _add_run(paragraph, text, bold=False, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = style.FONT_NAME
    run.font.size = size or style.BODY_SIZE
    return run


def render_docx(letter: dict, meta: dict, out_dir: Path) -> Path:
    """DOCX counterpart — same letterhead treatment as resume.docx, then a
    left-aligned business-letter body."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = style.MARGIN
    section.top_margin = section.bottom_margin = style.MARGIN

    normal = doc.styles["Normal"]
    normal.font.name = style.FONT_NAME
    normal.font.size = style.BODY_SIZE

    # Letterhead (mirrors render_docx.render_docx).
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(name_p, meta["name"], bold=True, size=style.NAME_SIZE)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        contact_p,
        f"{meta['location']} | {meta['phone']} | {meta['email']} | "
        f"{meta['linkedin']} | {meta['portfolio']}",
    )

    rule_p = doc.add_paragraph()
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), style.ACCENT_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()  # breathing room under the rule
    doc.add_paragraph(letter["date"])

    recipient = letter.get("recipient") or {}
    if recipient.get("addressee"):
        doc.add_paragraph(recipient["addressee"])
    if recipient.get("company"):
        company_line = recipient["company"]
        if recipient.get("location"):
            company_line += f", {recipient['location']}"
        doc.add_paragraph(company_line)

    doc.add_paragraph(letter["salutation"])
    for para in letter["paragraphs"]:
        doc.add_paragraph(para)

    doc.add_paragraph(letter["signoff"])
    signature = doc.add_paragraph()
    _add_run(signature, meta["name"], bold=True)

    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "cover_letter.docx"
    doc.save(str(docx_path))
    return docx_path
