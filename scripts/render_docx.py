"""instance.yaml -> resume.docx via python-docx.

Per TECH_SPEC.md §5: visually consistent with the PDF (same section order,
same heading hierarchy, same bullet content/order), not pixel-identical.
No live page-count check — the PDF's Tectonic count is the only thing that
gates the overflow loop.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "templates"))

from render_pdf import _order_experience  # noqa: E402
import resume_docx_style as style  # noqa: E402


def _set_bottom_border(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), style.HEADING_BORDER_COLOR)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_edge_borders(paragraph, edges=("top", "bottom")) -> None:
    """Draws accent hairline rules on the given edges of a paragraph — the DOCX
    counterpart of the PDF impact line's `\\rule`s above and below."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "3")
        el.set(qn("w:color"), style.ACCENT_COLOR)
        pBdr.append(el)
    pPr.append(pBdr)


def _add_run(paragraph, text: str, bold=False, italic=False, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = style.FONT_NAME
    run.font.size = size or style.BODY_SIZE
    return run


def _add_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, text, bold=True, size=style.HEADING_SIZE)
    _set_bottom_border(p)
    return p


def _add_split_line(doc, left: str, right: str, left_bold=False, size=None):
    """A paragraph with `left` flush left and `right` flush right via a
    right-aligned tab stop, mirroring the PDF's `\\hfill` layout."""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(style.USABLE_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    _add_run(p, left, bold=left_bold, size=size)
    _add_run(p, "\t" + right, size=size)
    return p


def _render_experience_entry(doc, exp: dict) -> None:
    end = exp["end"] if exp.get("end") else "Present"
    _add_split_line(doc, exp["title"], f"{exp['start']} - {end}", left_bold=True)

    company_p = doc.add_paragraph()
    company_p.paragraph_format.tab_stops.add_tab_stop(
        style.USABLE_WIDTH, WD_TAB_ALIGNMENT.RIGHT
    )
    _add_run(company_p, exp["company"])
    if exp.get("multinational") and exp.get("multinational_note"):
        _add_run(company_p, f" ({exp['multinational_note']})", italic=True, size=style.SMALL_SIZE)
    _add_run(company_p, "\t" + exp["location"])

    for bullet in exp["bullets"]:
        bp = doc.add_paragraph(style="List Bullet")
        _add_run(bp, bullet["text"])


def render_docx(instance: dict, out_dir: Path) -> Path:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = style.MARGIN
    section.right_margin = style.MARGIN
    section.top_margin = style.MARGIN
    section.bottom_margin = style.MARGIN

    normal = doc.styles["Normal"]
    normal.font.name = style.FONT_NAME
    normal.font.size = style.BODY_SIZE

    meta = instance["meta"]

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(name_p, meta["name"], bold=True, size=style.NAME_SIZE)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_line = f"{meta['location']} | {meta['phone']} | {meta['email']} | {meta['linkedin']} | {meta['github']}"
    _add_run(contact_p, contact_line)

    if instance.get("highlights"):
        impact_p = doc.add_paragraph()
        impact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        impact_text = "   ·   ".join(
            f"{h['value']} {h['label']}" for h in instance["highlights"]
        )
        impact_run = _add_run(impact_p, impact_text, bold=True)
        impact_run.font.color.rgb = RGBColor.from_string(style.ACCENT_COLOR)
        _set_edge_borders(impact_p, edges=("top", "bottom"))

    _add_heading(doc, "Summary")
    doc.add_paragraph(instance["summary"])

    ordered = _order_experience(instance["experience"])
    core_roles = [e for e in ordered if not e.get("additional")]
    additional_roles = [e for e in ordered if e.get("additional")]

    _add_heading(doc, "Experience")
    for exp in core_roles:
        _render_experience_entry(doc, exp)

    if additional_roles:
        _add_heading(doc, "Additional Experience")
        for exp in additional_roles:
            _render_experience_entry(doc, exp)

    _add_heading(doc, "Education")
    for edu in instance["education"]:
        _add_split_line(doc, edu["institution"], f"{edu['start']} - {edu['end']}", left_bold=True)
        _add_split_line(doc, edu["degree"], edu["location"])
        if edu.get("detail"):
            detail_p = doc.add_paragraph()
            _add_run(detail_p, edu["detail"], size=style.SMALL_SIZE)

    _add_heading(doc, "Skills")
    for group in instance["skills"]:
        p = doc.add_paragraph()
        _add_run(p, f"{group['label']}: ", bold=True)
        _add_run(p, ", ".join(group["items"]))

    if instance.get("certifications"):
        cert_p = doc.add_paragraph()
        _add_run(cert_p, "Certifications: ", bold=True)
        _add_run(cert_p, ", ".join(c["name"] for c in instance["certifications"]))

    _add_heading(doc, "Languages")
    languages_p = doc.add_paragraph()
    _add_run(
        languages_p,
        "    ".join(f"{lang['name']} ({lang['level']})" for lang in instance["languages"]),
    )

    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "resume.docx"
    doc.save(str(docx_path))
    return docx_path
